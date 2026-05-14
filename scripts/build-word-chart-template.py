#!/usr/bin/env python3
"""
Build a polished Word .docx DocGen template with {#ChartBucket} chart tags.

Writes docs/SurveyChartExample.docx. The visual goal matches the HTML demo —
cover page with stat cards, TOC, per-question chart pages with proper bar
charts, executive summary — using Word's own primitives:
  - Real fonts (Calibri/Helvetica), proper paragraph + run formatting
  - Colored eyebrow text + large blue titles + yellow accent dividers
  - Bar chart cells use Word percentage column widths (w:tcW w:type='pct')
    so the bar grows/shrinks with {percent_int}
  - Cycled per-bucket colors via {color_hex} merge tag — chart resolver
    emits the raw hex (no '#') for use in w:shd w:fill attributes
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Color palette — matches the HTML demo
BLUE = '1e3a8a'
GOLD = 'fbbf24'
TEXT = '1f2937'
MUTED = '6b7280'
LIGHT_MUTED = '9ca3af'
PANEL = 'f3f4f6'
PANEL_BLUE = 'eff6ff'
PANEL_BLUE_BORDER = 'c7d2fe'
TRACK = 'f3f4f6'
DEFAULT_FONT = 'Helvetica'

doc = Document()

# ---- Page setup ----
for section in doc.sections:
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

# Tight default paragraph spacing — Word's default 8pt-after looks flabby
styles = doc.styles
normal = styles['Normal']
normal.font.name = DEFAULT_FONT
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(TEXT)
normal.paragraph_format.space_after = Pt(0)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.line_spacing = 1.3

# ---------- Helpers ----------
def add_run(p, text, *, bold=False, size=None, color=None, font=DEFAULT_FONT):
    r = p.add_run(text)
    r.font.name = font
    if bold:
        r.bold = True
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    # Force eastAsia + ascii font for full coverage
    rPr = r._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), font)
    rFonts.set(qn('w:hAnsi'), font)
    return r

def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)

def set_cell_bg_via_token(cell, token):
    """Cell background fill where the value is a merge tag (e.g. {color_hex})."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), token)
    tc_pr.append(shd)

def set_cell_width_pct(cell, pct_or_token):
    """Cell width as a percentage. Numeric → multiplied by 50 (Word stores in
    50ths of a percent). Merge tag → appends '00' so {percent_int}=60 becomes
    '6000' meaning 60.00%."""
    tc_pr = cell._tc.get_or_add_tcPr()
    for ex in tc_pr.findall(qn('w:tcW')):
        tc_pr.remove(ex)
    tcw = OxmlElement('w:tcW')
    if isinstance(pct_or_token, str) and '{' in pct_or_token:
        tcw.set(qn('w:w'), pct_or_token + '00')
    else:
        tcw.set(qn('w:w'), str(int(pct_or_token) * 50))
    tcw.set(qn('w:type'), 'pct')
    tc_pr.append(tcw)

def set_cell_borders(cell, *, top=None, bottom=None, left=None, right=None, color='e5e7eb', sz=4):
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn('w:tcBorders'))
    if existing is not None:
        tc_pr.remove(existing)
    borders = OxmlElement('w:tcBorders')
    for side, val in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
        if val is None:
            continue
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), val)
        b.set(qn('w:sz'), str(sz))
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        borders.append(b)
    tc_pr.append(borders)

def remove_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn('w:tcBorders'))
    if existing is not None:
        tc_pr.remove(existing)
    borders = OxmlElement('w:tcBorders')
    for side in ('top', 'bottom', 'left', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'nil')
        borders.append(b)
    tc_pr.append(borders)

def remove_all_table_borders(table):
    tbl_pr = table._tbl.tblPr
    existing = tbl_pr.find(qn('w:tblBorders'))
    if existing is not None:
        tbl_pr.remove(existing)
    borders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'nil')
        borders.append(b)
    tbl_pr.append(borders)

def set_paragraph_spacing(p, *, before=None, after=None):
    pPr = p._p.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    if before is not None:
        spacing.set(qn('w:before'), str(before))
    if after is not None:
        spacing.set(qn('w:after'), str(after))

def page_break_before(p):
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement('w:pageBreakBefore')
    pPr.insert(0, pb)

def accent_divider(width_pct=15, color=GOLD, height_pt=3):
    """Insert a short colored bar — used as section divider, like the gold
    underline in the HTML demo."""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = False
    remove_all_table_borders(tbl)
    cell = tbl.rows[0].cells[0]
    set_cell_width_pct(cell, width_pct)
    remove_cell_borders(cell)
    set_cell_bg(cell, color)
    # Make the cell visually a thin bar — use tiny paragraph height
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, before=0, after=0)
    add_run(p, '', size=1)
    # Right cell empty filler
    spacer = tbl.rows[0].cells[1]
    set_cell_width_pct(spacer, 100 - width_pct)
    remove_cell_borders(spacer)
    return tbl

# ============================================================
# COVER PAGE
# ============================================================
# Vertical fill — push content down with spacer paragraphs
for _ in range(5):
    p = doc.add_paragraph()
    add_run(p, '', size=12)

# Eyebrow
p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
add_run(p, 'SURVEY RESPONSE REPORT', bold=True, size=11, color=GOLD)
set_paragraph_spacing(p, after=240)

# Title
p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
add_run(p, '{Name}', bold=True, size=36, color=BLUE)
set_paragraph_spacing(p, after=120)

# Subtitle
p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
add_run(p, 'Aggregated response analysis', size=13, color=MUTED)
set_paragraph_spacing(p, after=720)

# Stat cards — 3 cells, blue panel
stat_tbl = doc.add_table(rows=1, cols=3)
stat_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
stat_tbl.autofit = False
stat_tbl.width = Inches(6.0)
remove_all_table_borders(stat_tbl)
for cell, (number_tag, label) in zip(
    stat_tbl.rows[0].cells,
    [('{COUNT:Survey_Questions__r}', 'QUESTIONS'),
     ('{COUNT:Survey_Responses__r}', 'TOTAL RESPONSES'),
     ('{today:MMM d}', 'GENERATED')]
):
    set_cell_width_pct(cell, 33)
    set_cell_bg(cell, PANEL_BLUE)
    set_cell_borders(cell, top='single', bottom='single', left='single', right='single', color=PANEL_BLUE_BORDER, sz=6)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # First paragraph: large number
    np = cell.paragraphs[0]
    np.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_paragraph_spacing(np, before=160, after=80)
    add_run(np, number_tag, bold=True, size=26, color=BLUE)
    # Second paragraph: small caption
    lp = cell.add_paragraph()
    lp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_paragraph_spacing(lp, before=0, after=160)
    add_run(lp, label, bold=True, size=9, color=MUTED)

# Footer meta
p = doc.add_paragraph()
add_run(p, '', size=12)
set_paragraph_spacing(p, before=720)

p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
add_run(p, 'Prepared by {RunningUser.Name}', size=10, color=LIGHT_MUTED)

p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
add_run(p, '{today:MMMM d, yyyy}', size=10, color=LIGHT_MUTED)

# ============================================================
# TABLE OF CONTENTS
# ============================================================
p = doc.add_paragraph()
page_break_before(p)
add_run(p, 'Contents', bold=True, size=22, color=BLUE)
set_paragraph_spacing(p, after=120)

accent_divider(width_pct=12)

p = doc.add_paragraph()
add_run(p, '', size=8)

# Loop opener
p = doc.add_paragraph()
add_run(p, '{#Survey_Questions__r}')
set_paragraph_spacing(p, after=0)

# Per-iteration: a 2-col table — order number (gold) + question text
toc_tbl = doc.add_table(rows=1, cols=2)
toc_tbl.autofit = False
remove_all_table_borders(toc_tbl)
toc_row = toc_tbl.rows[0]
set_cell_width_pct(toc_row.cells[0], 8)
set_cell_width_pct(toc_row.cells[1], 92)
set_cell_borders(toc_row.cells[0], bottom='dotted', color='d1d5db', sz=4)
set_cell_borders(toc_row.cells[1], bottom='dotted', color='d1d5db', sz=4)
np = toc_row.cells[0].paragraphs[0]
np.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
set_paragraph_spacing(np, before=80, after=80)
add_run(np, '{Display_Order__c}.', bold=True, size=11, color=GOLD)
qp = toc_row.cells[1].paragraphs[0]
set_paragraph_spacing(qp, before=80, after=80)
add_run(qp, '{Question_Text__c}', size=11, color=TEXT)

p = doc.add_paragraph()
add_run(p, '{/Survey_Questions__r}')
set_paragraph_spacing(p, after=0)

# ============================================================
# PER-QUESTION CHART PAGES
# ============================================================
p = doc.add_paragraph()
page_break_before(p)
add_run(p, '{#Survey_Questions__r}')
set_paragraph_spacing(p, after=0)

# Eyebrow
p = doc.add_paragraph()
add_run(p, 'QUESTION {Display_Order__c}', bold=True, size=9, color=GOLD)
set_paragraph_spacing(p, after=60)

# Question title
p = doc.add_paragraph()
add_run(p, '{Question_Text__c}', bold=True, size=18, color=BLUE)
set_paragraph_spacing(p, after=160)

# Accent divider (gold short bar)
accent_divider(width_pct=10)

p = doc.add_paragraph()
add_run(p, '', size=8)

# Chart loop
p = doc.add_paragraph()
add_run(p, '{#ChartBucket:Survey_Responses__r:Selected_Answer__c}')
set_paragraph_spacing(p, after=0)

# Bar row: 4 cells — label | track-with-fill | spacer | figures
bar_tbl = doc.add_table(rows=1, cols=4)
bar_tbl.autofit = False
remove_all_table_borders(bar_tbl)
bar_row = bar_tbl.rows[0]
set_cell_width_pct(bar_row.cells[0], 28)
set_cell_width_pct(bar_row.cells[1], '{percent_int}')  # dynamic bar width
set_cell_width_pct(bar_row.cells[2], 50)
set_cell_width_pct(bar_row.cells[3], 22)

# Cell 0 — label
lp = bar_row.cells[0].paragraphs[0]
set_paragraph_spacing(lp, before=80, after=80)
add_run(lp, '{key_label}', bold=True, size=10, color=TEXT)
set_cell_borders(bar_row.cells[0], bottom='single', color='e5e7eb', sz=4)

# Cell 1 — colored bar (uses {color_hex} for per-bucket cycled fill)
bar_cell = bar_row.cells[1]
set_cell_bg_via_token(bar_cell, '{color_hex}')
set_cell_borders(bar_cell, bottom='single', color='e5e7eb', sz=4)
bp = bar_cell.paragraphs[0]
set_paragraph_spacing(bp, before=120, after=120)
add_run(bp, '', size=1)

# Cell 2 — spacer (light track)
spacer = bar_row.cells[2]
set_cell_bg(spacer, TRACK)
set_cell_borders(spacer, bottom='single', color='e5e7eb', sz=4)
sp = spacer.paragraphs[0]
set_paragraph_spacing(sp, before=120, after=120)
add_run(sp, '', size=1)

# Cell 3 — figures
figcell = bar_row.cells[3]
fp = figcell.paragraphs[0]
fp.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
set_paragraph_spacing(fp, before=80, after=80)
add_run(fp, '{count} ({percent}%)', size=10, color=MUTED)
set_cell_borders(figcell, bottom='single', color='e5e7eb', sz=4)

# Else branch
p = doc.add_paragraph()
add_run(p, '{:else}')
set_paragraph_spacing(p, after=0)

# Empty-state styled box
empty_tbl = doc.add_table(rows=1, cols=1)
empty_tbl.autofit = False
remove_all_table_borders(empty_tbl)
ecell = empty_tbl.rows[0].cells[0]
set_cell_width_pct(ecell, 100)
set_cell_bg(ecell, 'fef3c7')
set_cell_borders(ecell, left='single', color=GOLD, sz=12)
ep = ecell.paragraphs[0]
set_paragraph_spacing(ep, before=120, after=120)
add_run(ep, 'No responses recorded for this question yet.', size=10, color='92400e')

p = doc.add_paragraph()
add_run(p, '{/ChartBucket}')
set_paragraph_spacing(p, after=0)

# End question, page break before next iteration
loop_close = doc.add_paragraph()
page_break_before(loop_close)
add_run(loop_close, '{/Survey_Questions__r}')
set_paragraph_spacing(loop_close, after=0)

# ============================================================
# EXECUTIVE SUMMARY
# ============================================================
p = doc.add_paragraph()
page_break_before(p)
add_run(p, 'Methodology', bold=True, size=22, color=BLUE)
set_paragraph_spacing(p, after=120)

accent_divider(width_pct=12)

p = doc.add_paragraph()
add_run(p, '', size=8)

# Three styled blocks — table with shaded background + blue left border
for heading, body in [
    ('Aggregation.',
     ' Every chart in this report aggregates the underlying Survey Response records via DocGen\'s bucket-aggregation tag. The aggregation is performed server-side via SOQL GROUP BY, so the report renders at the same speed against 100 responses as it does against 100,000.'),
    ('Sort & null handling.',
     ' Buckets are sorted descending by count, with alphabetical ordering for ties. Null or blank answers collapse into a single "Not Specified" bucket so non-response rates are visible at a glance.'),
    ('Field-level security.',
     ' All queries run in USER_MODE — the values reflect what the running user is authorized to read. A user without access to a particular response field would see that question\'s chart fall back to an empty state, never a leaked aggregate.'),
]:
    btbl = doc.add_table(rows=1, cols=1)
    btbl.autofit = False
    remove_all_table_borders(btbl)
    bcell = btbl.rows[0].cells[0]
    set_cell_width_pct(bcell, 100)
    set_cell_bg(bcell, PANEL)
    set_cell_borders(bcell, left='single', color=BLUE, sz=12)
    bp = bcell.paragraphs[0]
    set_paragraph_spacing(bp, before=140, after=140)
    add_run(bp, heading, bold=True, size=11, color=BLUE)
    add_run(bp, body, size=11, color=TEXT)
    # Spacer paragraph after each block
    sp = doc.add_paragraph()
    add_run(sp, '', size=6)

# Footnote
fnp = doc.add_paragraph()
add_run(fnp, '', size=12)
set_paragraph_spacing(fnp, before=240)

ftbl = doc.add_table(rows=1, cols=1)
ftbl.autofit = False
remove_all_table_borders(ftbl)
fcell = ftbl.rows[0].cells[0]
set_cell_width_pct(fcell, 100)
set_cell_bg(fcell, 'f9fafb')
set_cell_borders(fcell, top='single', bottom='single', left='single', right='single', color='e5e7eb', sz=4)
fp = fcell.paragraphs[0]
set_paragraph_spacing(fp, before=140, after=140)
add_run(fp, 'Report generated by DocGen v1.91 on {today:MMMM d, yyyy} at {now:h:mm a}. Source: Survey {Name} containing {COUNT:Survey_Questions__r} survey questions and {COUNT:Survey_Responses__r} aggregated responses.', size=9, color=MUTED)

# ---- Save ----
out_path = '/Users/davemoudy/Desktop/Projects/DocGen/docs/SurveyChartExample.docx'
doc.save(out_path)
print(f"Wrote {out_path}")
print(f"Size: {os.path.getsize(out_path)} bytes")
